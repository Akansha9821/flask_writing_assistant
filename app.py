import base64
import io
import os
import re
import sqlite3
import uuid
from functools import wraps
from pathlib import Path
from datetime import datetime, timedelta

import pandas as pd
from docx import Document
from flask import (
    Flask,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Image as PDFImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    from deep_translator import GoogleTranslator
except ImportError:
    GoogleTranslator = None

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
UPLOAD_DIR = BASE_DIR / "uploads"
DB_PATH = DATA_DIR / "rpps_writer.db"

ALLOWED_EXTENSIONS = {"pdf", "png", "jpg", "jpeg", "doc", "docx", "txt", "csv", "xlsx"}
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg"}

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "replace-this-development-secret")
app.config.update(
    MAX_CONTENT_LENGTH=10 * 1024 * 1024,
    UPLOAD_FOLDER=str(UPLOAD_DIR),
    PERMANENT_SESSION_LIFETIME=timedelta(days=3650),
    SESSION_REFRESH_EACH_REQUEST=True,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=os.getenv("SESSION_COOKIE_SECURE", "false").lower() == "true",
)

DATA_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)

WRITING_TYPES = {
    "email": "Email",
    "letter": "Letter",
    "application": "Application",
    "mom": "Minutes of Meeting (MOM)",
}

CATEGORY_TREE = {
    "Office & Employment": {
        "leave_request": "Leave Request",
        "sick_leave": "Sick Leave Request",
        "office_followup": "Office Follow-up",
        "peer_communication": "Communication with Peers",
        "manager_request": "Manager Request",
        "hr_request": "HR Request",
        "work_from_home": "Work From Home Request",
        "attendance_correction": "Attendance Correction Request",
        "shift_change": "Shift Change Request",
        "promotion_request": "Promotion Request",
        "salary_increment": "Salary Increment Request",
        "transfer_request": "Transfer Request",
        "resignation": "Resignation Application",
        "notice_period_negotiation": "Notice Period Negotiation",
        "experience_letter": "Experience Letter Request",
        "relieving_letter": "Relieving Letter Request",
        "employment_verification": "Employment Verification Request",
        "employee_appreciation": "Employee Appreciation",
        "workplace_complaint": "Workplace Complaint",
    },
    "Recruitment & Job Communication": {
        "job_description_request": "Request Job Description",
        "job_interest_confirmation": "Confirm Interest in Job",
        "resume_submission": "Share Resume or Portfolio",
        "job_application": "Job Application",
        "internship_application": "Internship Application",
        "recruiter_followup": "Recruiter Follow-up",
        "application_status_followup": "Application Status Follow-up",
        "interview_availability": "Share Interview Availability",
        "interview_confirmation": "Interview Confirmation",
        "interview_rescheduling": "Interview Rescheduling Request",
        "interview_cancellation": "Interview Cancellation",
        "post_interview_thank_you": "Post-interview Thank-you",
        "interview_feedback": "Interview Feedback Request",
        "next_round_followup": "Next-round Follow-up",
        "salary_discussion": "Salary Discussion",
        "salary_negotiation": "Salary Negotiation",
        "document_submission": "Candidate Document Submission",
        "background_verification": "Background Verification",
        "reference_submission": "Professional Reference Submission",
        "offer_letter_followup": "Offer Letter Follow-up",
        "offer_acceptance": "Offer Acceptance",
        "offer_rejection": "Offer Rejection",
        "joining_confirmation": "Joining Date Confirmation",
        "joining_extension": "Joining Date Extension Request",
        "application_withdrawal": "Job Application Withdrawal",
        "referral_request": "Job Referral Request",
    },
   "Banking & Finance": {
        "bank_service": "Bank Customer Service",
        "account_opening": "Bank Account Opening Request",
        "account_statement": "Bank Statement Request",
        "payment_issue": "Payment Issue",
        "payment_confirmation": "Payment Confirmation",
        "refund_request": "Refund Request",
        "refund_followup": "Refund Follow-up",
        "transaction_dispute": "Transaction Dispute",
        "unauthorized_transaction": "Unauthorized Transaction Report",
        "failed_transaction": "Failed Transaction Complaint",
        "account_update": "Account Update Request",
        "loan_request": "Loan-related Request",
        "loan_application": "Loan Application",
        "loan_status_followup": "Loan Status Follow-up",
        "credit_debit_card_issue": "Credit or Debit Card Issue",
        "emi_request": "EMI-related Request",
        "insurance_claim": "Insurance Claim",
        "invoice_submission": "Invoice Submission",
        "invoice_correction": "Invoice Correction",
        "payment_reminder": "Payment Reminder",
        "tax_document_request": "Tax Document Request",
    },
    "Customer & Product Service": {
        "product_enquiry": "Product Enquiry",
        "service_enquiry": "Service Enquiry",
        "device_complaint": "Device Complaint",
        "customer_service": "Customer Service Request",
        "technical_support": "Technical Support Request",
        "product_complaint": "Product Complaint",
        "service_complaint": "Service Complaint",
        "complaint_escalation": "Complaint Escalation",
        "replacement_request": "Replacement Request",
        "refund_request": "Refund Request",
        "cancellation_request": "Cancellation Request",
        "warranty_request": "Warranty Request",
        "delivery_delay": "Delivery Delay Complaint",
        "order_status": "Order Status Request",
        "account_access_issue": "Account Access Issue",
        "service_followup": "Service Follow-up",
        "service_feedback": "Product or Service Feedback",
    },
    "Study & University": {
        "admission_enquiry": "Admission Enquiry",
        "course_enquiry": "Course Enquiry",
        "exam": "Exam Related",
        "exam_leave": "Exam Leave Request",
        "exam_schedule": "Exam Schedule Enquiry",
        "project": "Project Related",
        "project_submission": "Project Submission",
        "project_guide_followup": "Project Guide Follow-up",
        "placement": "Placement Related",
        "practical": "Practical Related",
        "practical_rescheduling": "Practical Rescheduling Request",
        "viva": "Viva Related",
        "viva_rescheduling": "Viva Rescheduling Request",
        "assignment_submission": "Assignment Submission",
        "assignment_extension": "Assignment Extension Request",
        "fees_payment": "Fees or Payment Related",
        "class": "Class Related",
        "attendance": "Attendance Related",
        "attendance_correction": "Attendance Correction Request",
        "academic_complaint": "Academic Complaint",
        "certificate": "Certificate Request",
        "result_correction": "Result Correction Request",
        "scholarship": "Scholarship Application",
        "recommendation_letter": "Recommendation Letter Request",
        "internship_permission": "Internship Permission Request",
        "research_proposal": "Research Proposal Submission",
    },
     "Business & Corporate": {
        "business_proposal": "Business Proposal",
        "partnership_proposal": "Partnership Proposal",
        "client_introduction": "Client Introduction",
        "project_update": "Project Status Update",
        "project_delay": "Project Delay Notification",
        "business_followup": "Business Follow-up",
        "quotation_request": "Quotation Request",
        "quotation_submission": "Quotation Submission",
        "contract_discussion": "Contract Discussion",
        "vendor_communication": "Vendor Communication",
        "customer_followup": "Customer Follow-up",
        "business_invitation": "Business Invitation",
        "company_announcement": "Company Announcement",
        "mou_request": "Memorandum of Understanding Request",
    },
     "Meeting & Administration": {
        "meeting_request": "Meeting Request",
        "meeting_invitation": "Meeting Invitation",
        "meeting_confirmation": "Meeting Confirmation",
        "meeting_rescheduling": "Meeting Rescheduling Request",
        "meeting_cancellation": "Meeting Cancellation",
        "meeting_agenda": "Meeting Agenda",
        "mom_internal": "Internal Meeting MOM",
        "mom_client": "Client Meeting MOM",
        "mom_project": "Project Meeting MOM",
        "mom_review": "Review Meeting MOM",
        "action_item_summary": "Action-item Summary",
        "decision_summary": "Decision Summary",
        "meeting_followup": "Meeting Follow-up",
        "participant_reminder": "Meeting Participant Reminder",
        "general_application": "General Application",
        "permission_application": "Permission Application",
        "complaint_application": "Complaint Application",
    },
     "Government & Legal": {
        "government_application": "Government Application",
        "information_request": "Official Information Request",
        "certificate_application": "Certificate Application",
        "license_permit_request": "Licence or Permit Request",
        "police_complaint": "Police Complaint",
        "legal_notice": "Legal Notice",
        "legal_response": "Response to Legal Notice",
        "consumer_complaint": "Consumer Complaint",
        "rti_application": "RTI Application",
        "municipal_complaint": "Municipal Complaint",
        "document_correction": "Official Document Correction Request",
        "grievance_submission": "Official Grievance Submission",
    },"Healthcare": {
        "appointment_request": "Medical Appointment Request",
        "appointment_rescheduling": "Appointment Rescheduling Request",
        "medical_leave": "Medical Leave Request",
        "medical_report": "Medical Report Request",
        "prescription_request": "Prescription Request",
        "insurance_approval": "Medical Insurance Approval Request",
        "hospital_complaint": "Hospital Complaint",
        "billing_enquiry": "Medical Billing Enquiry",
        "test_result_followup": "Medical Test Result Follow-up",
        "medical_certificate": "Medical Certificate Request",
    },

    "Sales & Marketing": {
        "sales_introduction": "Sales Introduction",
        "product_promotion": "Product Promotion",
        "lead_followup": "Sales Lead Follow-up",
        "customer_outreach": "Customer Outreach",
        "marketing_proposal": "Marketing Proposal",
        "campaign_approval": "Marketing Campaign Approval",
        "sponsorship_request": "Sponsorship Request",
        "collaboration_request": "Collaboration Request",
        "event_promotion": "Event Promotion",
        "customer_feedback": "Customer Feedback Request",
        "subscription_renewal": "Subscription Renewal",
        "discount_offer": "Discount Offer",
    },

    "IT & Technical Communication": {
        "bug_report": "Bug Report",
        "access_request": "System Access Request",
        "password_reset": "Password Reset Request",
        "software_installation": "Software Installation Request",
        "system_downtime": "System Downtime Notification",
        "security_incident": "Security Incident Report",
        "data_request": "Data Request",
        "feature_request": "Software Feature Request",
        "deployment_approval": "Deployment Approval Request",
        "technical_followup": "Technical Issue Follow-up",
        "account_deactivation": "Account Deactivation Request",
    },

    "Notices & Announcements": {
        "general_notice": "General Notice",
        "holiday_notice": "Holiday Notice",
        "event_announcement": "Event Announcement",
        "policy_update": "Policy Update",
        "maintenance_notice": "Maintenance Notice",
        "meeting_notice": "Meeting Notice",
        "examination_notice": "Examination Notice",
        "office_closure": "Office Closure Notice",
        "emergency_announcement": "Emergency Announcement",
        "deadline_reminder": "Deadline Reminder",
    },

    "Personal Communication": {
        "personal_invitation": "Personal Invitation",
        "thank_you": "Thank-you Message",
        "apology": "Apology Message",
        "congratulations": "Congratulations Message",
        "condolence": "Condolence Message",
        "birthday_wish": "Birthday Wish",
        "festival_greeting": "Festival Greeting",
        "personal_request": "Personal Request",
        "friendly_followup": "Friendly Follow-up",
        "family_communication": "Family Communication",
    },

    "General": {
        "help": "General Help",
        "complaint": "General Complaint",
        "request": "General Request",
        "enquiry": "General Enquiry",
        "followup": "General Follow-up",
        "invitation": "Invitation",
        "thank_you": "Thank-you Message",
        "apology": "Apology Message",
        "feedback": "General Feedback",
    },
}

LANGUAGES = {
    "en": "English",
    "hi": "Hindi",
    "bn": "Bengali",
    "mr": "Marathi",
    "gu": "Gujarati",
    "pa": "Punjabi",
    "ta": "Tamil",
    "te": "Telugu",
    "kn": "Kannada",
    "ml": "Malayalam",
    "ur": "Urdu",
    "ne": "Nepali",
    "fr": "French",
    "de": "German",
    "es": "Spanish",
    "it": "Italian",
    "pt": "Portuguese",
    "ar": "Arabic",
    "ja": "Japanese",
    "ko": "Korean",
    "zh-CN": "Chinese (Simplified)",
    "ru": "Russian",
}

INFO_PAGES = {
    "about": {
        "eyebrow": "Who we are",
        "title": "About RPPS",
        "icon": "building",
        "intro": "RPPS Writing Assistant helps people turn everyday ideas into clear, polished professional communication.",
        "sections": [
            {
                "heading": "Our purpose",
                "body": "We make formal writing easier for office, academic, banking, customer-service and personal communication needs.",
            },
            {
                "heading": "How it works",
                "body": "Choose a writing type, provide the important details and review the structured draft before printing or exporting it.",
            },
            {
                "heading": "Built for practical work",
                "body": "The workspace keeps recent writings together and supports PDF, document and image exports for convenient sharing.",
            },
            {
                "heading": "Responsible assistance",
                "body": "Every generated draft should be reviewed by the user to confirm names, facts, tone and context before it is sent.",
            },
        ],
    },
    "services": {
        "eyebrow": "What we offer",
        "title": "Services",
        "icon": "grid-1x2",
        "intro": "A focused set of tools for drafting, refining, signing and exporting professional writing.",
        "sections": [
            {
                "heading": "Email and letter drafting",
                "body": "Create formal emails, informal emails, applications, letters and meeting minutes from guided inputs.",
            },
            {
                "heading": "Writing organization",
                "body": "Access recent work and history from your dashboard so drafts remain easy to find and reuse.",
            },
            {
                "heading": "Signatures and attachments",
                "body": "Add a drawn or uploaded signature and keep a supporting attachment associated with your writing.",
            },
            {
                "heading": "Flexible exports",
                "body": "Print a clean layout or export supported writings as PDF, DOCX, PNG or text files.",
            },
        ],
    },
    "contact": {
        "eyebrow": "Get in touch",
        "title": "Contact",
        "icon": "envelope",
        "intro": "Questions, feedback or technical concerns are welcome. Use the contact option below to reach the RPPS team.",
        "sections": [
            {
                "heading": "General enquiries",
                "body": "Contact us with questions about the writing assistant, its features or how to get started.",
                "link": {"href": "mailto:admin@rppssoftware.com", "label": "Email RPPS"},
            },
            {
                "heading": "Technical support",
                "body": "Include a short description of the issue, the page where it occurred and any error message you received.",
                "link": {"href": "/help", "label": "Visit Help and Support"},
            },
            {
                "heading": "Product feedback",
                "body": "Tell us which workflows are useful and what would make professional writing easier for you.",
                "link": {"href": "mailto:admin@rppssoftware.com?subject=Writing%20Assistant%20Feedback", "label": "Send feedback"},
            },
            {
                "heading": "Response information",
                "body": "Please avoid including passwords or other sensitive information in support messages.",
            },
        ],
    },
    "privacy": {
        "eyebrow": "Your information",
        "title": "Privacy Policy",
        "icon": "shield-check",
        "intro": "This page explains the categories of information used by the writing assistant and the choices available to users.",
        "sections": [
            {
                "heading": "Information you provide",
                "body": "The service stores account details and the writing, attachments and signatures you choose to submit so it can provide its features.",
            },
            {
                "heading": "How information is used",
                "body": "Information is used to authenticate your account, create and manage writings, produce exports and maintain service operation.",
            },
            {
                "heading": "Account security",
                "body": "Use a strong password, protect access to your device and sign out when using a shared computer.",
            },
            {
                "heading": "Privacy questions",
                "body": "Contact RPPS if you have a question about information associated with your account.",
                "link": {"href": "mailto:admin@rppssoftware.com?subject=Privacy%20Question", "label": "Ask a privacy question"},
            },
        ],
    },
    "terms": {
        "eyebrow": "Using the service",
        "title": "Terms and Conditions",
        "icon": "file-earmark-text",
        "intro": "These practical terms describe responsible use of RPPS Writing Assistant and its generated drafts.",
        "sections": [
            {
                "heading": "Review your drafts",
                "body": "Generated content is a drafting aid. You are responsible for reviewing accuracy, suitability and recipients before use.",
            },
            {
                "heading": "Acceptable use",
                "body": "Do not use the service to create unlawful, deceptive, abusive or rights-infringing material.",
            },
            {
                "heading": "Your account",
                "body": "Keep account credentials confidential and provide accurate registration information when using the service.",
            },
            {
                "heading": "Service availability",
                "body": "Features may be updated or temporarily unavailable for maintenance, security or operational reasons.",
            },
        ],
    },
    "help": {
        "eyebrow": "Guidance and support",
        "title": "Help and Support",
        "icon": "life-preserver",
        "intro": "Quick guidance for creating, reviewing, exporting and managing your writing.",
        "sections": [
            {
                "heading": "Create a writing",
                "body": "Open Create Writing, select the document type and category, add the essential details and submit the form to generate a draft.",
            },
            {
                "heading": "Review before exporting",
                "body": "Check the recipient, subject, dates, names and tone. Then use Print, PDF, DOCX or Image from the preview page.",
            },
            {
                "heading": "Find previous work",
                "body": "Use Dashboard for recent activity or History to browse the writings saved to your account.",
            },
            {
                "heading": "Still need help?",
                "body": "Send a concise description of the problem and include any displayed error message.",
                "link": {"href": "mailto:admin@rppssoftware.com?subject=Writing%20Assistant%20Support", "label": "Contact support"},
            },
        ],
    },
    "accessibility": {
        "eyebrow": "Inclusive experience",
        "title": "Accessibility",
        "icon": "universal-access",
        "intro": "RPPS Writing Assistant aims to provide a clear experience that works across common devices and input methods.",
        "sections": [
            {
                "heading": "Keyboard access",
                "body": "Interactive controls are designed with standard links, buttons and form fields that can be reached using a keyboard.",
            },
            {
                "heading": "Readable presentation",
                "body": "Layouts use clear headings, responsive spacing and strong text contrast to support comfortable reading.",
            },
            {
                "heading": "Responsive layouts",
                "body": "Pages adapt to desktop, tablet and mobile screens while keeping key actions available.",
            },
            {
                "heading": "Report a barrier",
                "body": "If something prevents you from using the service, tell us which page and assistive technology were involved.",
                "link": {"href": "mailto:admin@rppssoftware.com?subject=Accessibility%20Feedback", "label": "Send accessibility feedback"},
            },
        ],
    },
}


PREDEFINED_TITLES = {
    "Office & Employment": {
        "leave_request": "Request for Leave",
        "sick_leave": "Request for Sick Leave",
        "office_followup": "Follow-up Regarding Office Request",
        "peer_communication": "Regarding Team Communication",
        "manager_request": "Request for Manager's Assistance",
        "hr_request": "Request for HR Assistance",
        "work_from_home": "Request to Work from Home",
        "attendance_correction": "Request for Attendance Correction",
        "shift_change": "Request for Shift Change",
        "promotion_request": "Request for Promotion Consideration",
        "salary_increment": "Request for Salary Revision",
        "transfer_request": "Request for Department or Location Transfer",
        "resignation": "Formal Resignation from My Position",
        "notice_period_negotiation": "Request for Notice Period Revision",
        "experience_letter": "Request for Experience Letter",
        "relieving_letter": "Request for Relieving Letter",
        "employment_verification": "Request for Employment Verification",
        "employee_appreciation": "Appreciation for Outstanding Contribution",
        "workplace_complaint": "Formal Workplace Complaint",
    },

    "Recruitment & Job Communication": {
        "job_description_request": "Request for Complete Job Description",
        "job_interest_confirmation": "Confirmation of Interest in the Position",
        "resume_submission": "Application and Resume for [Job Title]",
        "job_application": "Application for the Position of [Job Title]",
        "internship_application": "Application for [Internship Title] Internship",
        "recruiter_followup": "Follow-up Regarding [Job Title] Opportunity",
        "application_status_followup": "Follow-up on Application for [Job Title]",
        "interview_availability": "Interview Availability for [Job Title]",
        "interview_confirmation": "Confirmation of Interview for [Job Title]",
        "interview_rescheduling": "Request to Reschedule Interview for [Job Title]",
        "interview_cancellation": "Interview Cancellation for [Job Title]",
        "post_interview_thank_you": "Thank You for the [Job Title] Interview",
        "interview_feedback": "Request for Interview Feedback",
        "next_round_followup": "Follow-up Regarding the Next Interview Round",
        "salary_discussion": "Discussion Regarding Compensation for [Job Title]",
        "salary_negotiation": "Request to Review the Compensation Offer",
        "document_submission": "Submission of Candidate Documents",
        "background_verification": "Documents for Background Verification",
        "reference_submission": "Submission of Professional References",
        "offer_letter_followup": "Follow-up Regarding Offer Letter for [Job Title]",
        "offer_acceptance": "Acceptance of Offer for [Job Title]",
        "offer_rejection": "Response to Employment Offer for [Job Title]",
        "joining_confirmation": "Confirmation of Joining Date",
        "joining_extension": "Request for Extension of Joining Date",
        "application_withdrawal": "Withdrawal of Application for [Job Title]",
        "referral_request": "Request for Referral for [Job Title]",
    },

    "Banking & Finance": {
        "bank_service": "Request for Banking Assistance",
        "account_opening": "Request to Open a Bank Account",
        "account_statement": "Request for Bank Account Statement",
        "payment_issue": "Assistance Required Regarding Payment Issue",
        "payment_confirmation": "Confirmation of Payment",
        "refund_request": "Request for Refund",
        "refund_followup": "Follow-up Regarding Pending Refund",
        "transaction_dispute": "Dispute Regarding Transaction [Transaction ID]",
        "unauthorized_transaction": "Urgent: Unauthorized Transaction Report",
        "failed_transaction": "Complaint Regarding Failed Transaction",
        "account_update": "Request to Update Account Information",
        "loan_request": "Enquiry Regarding Loan Services",
        "loan_application": "Application for [Loan Type] Loan",
        "loan_status_followup": "Follow-up on Loan Application",
        "credit_debit_card_issue": "Assistance Required for Card Issue",
        "emi_request": "Request Regarding EMI Facility",
        "insurance_claim": "Submission of Insurance Claim",
        "invoice_submission": "Submission of Invoice [Invoice Number]",
        "invoice_correction": "Request for Correction of Invoice",
        "payment_reminder": "Reminder for Outstanding Payment",
        "tax_document_request": "Request for Tax-related Documents",
    },

    "Customer & Product Service": {
        "product_enquiry": "Enquiry Regarding [Product Name]",
        "service_enquiry": "Enquiry Regarding [Service Name]",
        "device_complaint": "Complaint Regarding [Device Name]",
        "customer_service": "Request for Customer Service Assistance",
        "technical_support": "Technical Support Required for [Product Name]",
        "product_complaint": "Complaint Regarding [Product Name]",
        "service_complaint": "Complaint Regarding [Service Name]",
        "complaint_escalation": "Escalation of Unresolved Complaint",
        "replacement_request": "Request for Product Replacement",
        "refund_request": "Request for Product or Service Refund",
        "cancellation_request": "Request to Cancel [Order or Service]",
        "warranty_request": "Warranty Service Request for [Product Name]",
        "delivery_delay": "Complaint Regarding Delayed Delivery",
        "order_status": "Request for Order Status – [Order ID]",
        "account_access_issue": "Unable to Access Customer Account",
        "service_followup": "Follow-up Regarding Service Request",
        "service_feedback": "Feedback Regarding [Product or Service Name]",
    },

    "Study & University": {
        "admission_enquiry": "Enquiry Regarding Admission to [Course Name]",
        "course_enquiry": "Request for Information About [Course Name]",
        "exam": "Request Regarding [Subject Name] Examination",
        "exam_leave": "Request for Examination Leave",
        "exam_schedule": "Enquiry Regarding Examination Schedule",
        "project": "Request Regarding Academic Project",
        "project_submission": "Submission of [Project Title]",
        "project_guide_followup": "Follow-up Regarding Project Guidance",
        "placement": "Request Regarding Placement Process",
        "practical": "Request Regarding Practical Examination",
        "practical_rescheduling": "Request to Reschedule Practical Examination",
        "viva": "Request Regarding Final Viva-Voce",
        "viva_rescheduling": "Request to Reschedule Viva-Voce",
        "assignment_submission": "Submission of [Assignment Title]",
        "assignment_extension": "Request for Assignment Submission Extension",
        "fees_payment": "Request Regarding Fee Payment",
        "class": "Request Regarding [Subject Name] Class",
        "attendance": "Request Regarding Attendance Record",
        "attendance_correction": "Request for Attendance Correction",
        "academic_complaint": "Formal Academic Complaint",
        "certificate": "Request for Academic Certificate",
        "result_correction": "Request for Correction in Examination Result",
        "scholarship": "Application for [Scholarship Name]",
        "recommendation_letter": "Request for Letter of Recommendation",
        "internship_permission": "Request for Internship Permission",
        "research_proposal": "Submission of Research Proposal – [Title]",
    },

    "Business & Corporate": {
        "business_proposal": "Business Proposal for [Project or Service]",
        "partnership_proposal": "Proposal for Business Partnership",
        "client_introduction": "Introduction and Business Collaboration Opportunity",
        "project_update": "Project Status Update – [Project Name]",
        "project_delay": "Notification of Delay in [Project Name]",
        "business_followup": "Follow-up Regarding Business Discussion",
        "quotation_request": "Request for Quotation – [Product or Service]",
        "quotation_submission": "Submission of Quotation – [Quotation Number]",
        "contract_discussion": "Discussion Regarding [Contract Name]",
        "vendor_communication": "Regarding Vendor Services for [Project Name]",
        "customer_followup": "Follow-up Regarding [Product or Service]",
        "business_invitation": "Invitation to [Business Event or Meeting]",
        "company_announcement": "Company Announcement – [Subject]",
        "mou_request": "Proposal for Memorandum of Understanding",
    },

    "Meeting & Administration": {
        "meeting_request": "Request to Schedule a Meeting",
        "meeting_invitation": "Invitation: [Meeting Subject]",
        "meeting_confirmation": "Confirmation of Meeting on [Date]",
        "meeting_rescheduling": "Request to Reschedule [Meeting Subject]",
        "meeting_cancellation": "Cancellation of Meeting Scheduled for [Date]",
        "meeting_agenda": "Agenda for [Meeting Name]",
        "mom_internal": "Minutes of Internal Meeting – [Date]",
        "mom_client": "Minutes of Client Meeting – [Client Name]",
        "mom_project": "Project Meeting Minutes – [Project Name]",
        "mom_review": "Minutes of Review Meeting – [Date]",
        "action_item_summary": "Action Items from [Meeting Name]",
        "decision_summary": "Summary of Decisions from [Meeting Name]",
        "meeting_followup": "Follow-up on [Meeting Name]",
        "participant_reminder": "Reminder: [Meeting Name] on [Date]",
        "general_application": "Application Regarding [Subject]",
        "permission_application": "Request for Permission to [Purpose]",
        "complaint_application": "Formal Complaint Regarding [Issue]",
    },

    "Government & Legal": {
        "government_application": "Application Regarding [Service or Scheme]",
        "information_request": "Request for Official Information",
        "certificate_application": "Application for [Certificate Name]",
        "license_permit_request": "Application for [Licence or Permit Name]",
        "police_complaint": "Formal Complaint Regarding [Incident]",
        "legal_notice": "Legal Notice Regarding [Matter]",
        "legal_response": "Response to Legal Notice Dated [Date]",
        "consumer_complaint": "Consumer Complaint Regarding [Product or Service]",
        "rti_application": "Application Under the Right to Information Act",
        "municipal_complaint": "Complaint Regarding Municipal Service",
        "document_correction": "Request for Correction in Official Document",
        "grievance_submission": "Submission of Grievance Regarding [Issue]",
    },

    "Healthcare": {
        "appointment_request": "Request for Medical Appointment",
        "appointment_rescheduling": "Request to Reschedule Medical Appointment",
        "medical_leave": "Request for Medical Leave",
        "medical_report": "Request for Medical Report",
        "prescription_request": "Request for Prescription",
        "insurance_approval": "Request for Medical Insurance Approval",
        "hospital_complaint": "Complaint Regarding Hospital Service",
        "billing_enquiry": "Enquiry Regarding Medical Bill",
        "test_result_followup": "Follow-up Regarding Medical Test Results",
        "medical_certificate": "Request for Medical Certificate",
    },

    "Sales & Marketing": {
        "sales_introduction": "Introduction to [Product or Service]",
        "product_promotion": "Special Offer on [Product Name]",
        "lead_followup": "Follow-up Regarding [Product or Service]",
        "customer_outreach": "Discover How [Product Name] Can Help You",
        "marketing_proposal": "Marketing Proposal for [Campaign Name]",
        "campaign_approval": "Request for Approval of [Campaign Name]",
        "sponsorship_request": "Sponsorship Request for [Event or Project]",
        "collaboration_request": "Proposal for Marketing Collaboration",
        "event_promotion": "Invitation to [Event Name]",
        "customer_feedback": "Request for Your Valuable Feedback",
        "subscription_renewal": "Reminder to Renew Your Subscription",
        "discount_offer": "Special Discount Offer on [Product or Service]",
    },

    "IT & Technical Communication": {
        "bug_report": "Bug Report: [Brief Issue Description]",
        "access_request": "Request for Access to [System Name]",
        "password_reset": "Request to Reset Account Password",
        "software_installation": "Request to Install [Software Name]",
        "system_downtime": "System Downtime Notification",
        "security_incident": "Urgent: Security Incident Report",
        "data_request": "Request for [Data or Report Name]",
        "feature_request": "Feature Request: [Feature Name]",
        "deployment_approval": "Request for Deployment Approval",
        "technical_followup": "Follow-up Regarding Technical Issue",
        "account_deactivation": "Request for Account Deactivation",
    },

    "Notices & Announcements": {
        "general_notice": "Notice: [Subject]",
        "holiday_notice": "Holiday Notice for [Date or Occasion]",
        "event_announcement": "Announcement: [Event Name]",
        "policy_update": "Important Update to [Policy Name]",
        "maintenance_notice": "Scheduled Maintenance Notification",
        "meeting_notice": "Notice of Meeting on [Date]",
        "examination_notice": "Examination Notice – [Exam Name]",
        "office_closure": "Office Closure Notice",
        "emergency_announcement": "Urgent Announcement: [Subject]",
        "deadline_reminder": "Reminder: Deadline for [Task]",
    },

    "Personal Communication": {
        "personal_invitation": "Invitation to [Event Name]",
        "thank_you": "Thank You for [Reason]",
        "apology": "Sincere Apology Regarding [Matter]",
        "congratulations": "Congratulations on [Achievement]",
        "condolence": "Heartfelt Condolences",
        "birthday_wish": "Happy Birthday, [Name]!",
        "festival_greeting": "Warm Wishes for [Festival Name]",
        "personal_request": "Personal Request Regarding [Subject]",
        "friendly_followup": "Following Up Regarding [Subject]",
        "family_communication": "Regarding [Family Matter]",
    },

    "General": {
        "help": "Request for Assistance Regarding [Subject]",
        "complaint": "Complaint Regarding [Issue]",
        "request": "Request Regarding [Subject]",
        "enquiry": "Enquiry Regarding [Subject]",
        "followup": "Follow-up Regarding [Subject]",
        "invitation": "Invitation to [Event Name]",
        "thank_you": "Thank You for [Reason]",
        "apology": "Apology Regarding [Matter]",
        "feedback": "Feedback Regarding [Subject]",
    },
}

PREDEFINED_DESCRIPTIONS = {
    "Office & Employment": {
        "leave_request": (
            "I am writing to request leave from {start_date} to {end_date} "
            "due to {main_subject}. I kindly request you to approve my leave "
            "for the specified period. I will complete or hand over my pending "
            "responsibilities before the leave begins."
        ),
        "sick_leave": (
            "I am unable to attend work from {start_date} to {end_date} due to "
            "{main_subject}. I request you to kindly approve my sick leave for "
            "this period. I will resume my responsibilities as soon as I recover."
        ),
        "work_from_home": (
            "I am requesting permission to work from home from {start_date} to "
            "{end_date} due to {main_subject}. I will remain available during "
            "working hours and ensure that all assigned responsibilities are "
            "completed without interruption."
        ),
        "attendance_correction": (
            "I am writing to request correction of my attendance record for the "
            "period from {start_date} to {end_date}. The correction is required "
            "because {main_subject}. Kindly review the relevant records and update "
            "my attendance accordingly."
        ),
        "office_followup": (
            "I am following up on my request submitted on {start_date} regarding "
            "{main_subject}. As the matter remains pending, I would appreciate an "
            "update by {end_date}. Please let me know if any additional information "
            "is required from my side."
        ),
        "manager_request": (
            "I am writing to request your assistance regarding {main_subject}. "
            "This matter is relevant for the period from {start_date} to {end_date}. "
            "I would appreciate your guidance and approval to proceed further."
        ),
        "hr_request": (
            "I am writing to request HR assistance regarding {main_subject}. "
            "The matter concerns the period from {start_date} to {end_date}. "
            "Kindly review my request and advise me about the required next steps."
        ),
        "shift_change": (
            "I request a change in my assigned work shift from {start_date} to "
            "{end_date} due to {main_subject}. I will ensure that the requested "
            "change does not affect my responsibilities or the team's work."
        ),
        "promotion_request": (
            "I would like to request consideration for a promotion based on "
            "{main_subject}. My work and contributions from {start_date} to "
            "{end_date} demonstrate my readiness to take on additional "
            "responsibilities. I would appreciate an opportunity to discuss this."
        ),
        "salary_increment": (
            "I am writing to request a review of my compensation based on "
            "{main_subject}. During the period from {start_date} to {end_date}, "
            "I have consistently fulfilled my responsibilities and contributed "
            "to the organisation's objectives. I request a discussion regarding "
            "an appropriate salary revision."
        ),
        "transfer_request": (
            "I request a transfer for the period beginning {start_date}, preferably "
            "by {end_date}, due to {main_subject}. I will complete the necessary "
            "handover and support a smooth transition if the request is approved."
        ),
        "resignation": (
            "Please accept this message as formal notice of my resignation due to "
            "{main_subject}. I am submitting my resignation on {start_date}, and "
            "I request that {end_date} be considered my final working day. I will "
            "complete the required handover during the notice period."
        ),
        "notice_period_negotiation": (
            "I request a revision of my notice period from {start_date} to "
            "{end_date} due to {main_subject}. I will complete all essential "
            "handover activities and assist with the transition before my "
            "proposed last working date."
        ),
        "experience_letter": (
            "I request the issuance of my experience letter for my employment "
            "period from {start_date} to {end_date}. The document is required for "
            "{main_subject}. Kindly issue it at your earliest convenience."
        ),
        "relieving_letter": (
            "I request the issuance of my relieving letter for my employment "
            "period from {start_date} to {end_date}. The letter is required for "
            "{main_subject}. Please let me know if any formalities remain pending."
        ),
        "employment_verification": (
            "I request verification of my employment for the period from "
            "{start_date} to {end_date}. The verification is required for "
            "{main_subject}. Kindly provide the necessary confirmation or "
            "documentation."
        ),
        "workplace_complaint": (
            "I wish to formally report a workplace concern that occurred between "
            "{start_date} and {end_date}. The matter concerns {main_subject}. "
            "I request a confidential and impartial review and appropriate action "
            "according to organisational policy."
        ),
    },

    "Recruitment & Job Communication": {
        "job_description_request": (
            "I am interested in the opportunity discussed on {start_date} regarding "
            "{main_subject}. Please share the complete job description, required "
            "skills, responsibilities, work location, employment type and selection "
            "process by {end_date}, if possible."
        ),
        "job_interest_confirmation": (
            "I am writing to confirm my interest in the position concerning "
            "{main_subject}. Following our discussion on {start_date}, I would "
            "like to proceed with the recruitment process and remain available "
            "for the next steps until {end_date}."
        ),
        "resume_submission": (
            "Please find my resume and relevant professional details for the "
            "opportunity concerning {main_subject}. I am available for the "
            "recruitment process from {start_date} to {end_date}. I believe my "
            "experience and skills align with the requirements of the position."
        ),
        "job_application": (
            "I am applying for the position concerning {main_subject}. My relevant "
            "experience and skills make me a suitable candidate for this opportunity. "
            "I am available for interviews between {start_date} and {end_date} and "
            "would appreciate your consideration of my application."
        ),
        "internship_application": (
            "I am applying for the internship opportunity concerning {main_subject}. "
            "I am available from {start_date} to {end_date} and am interested in "
            "gaining practical experience while contributing my knowledge and skills "
            "to the organisation."
        ),
        "recruiter_followup": (
            "I am following up regarding {main_subject}, which we discussed on "
            "{start_date}. I remain interested in the opportunity and would "
            "appreciate an update on my application or the next selection stage "
            "by {end_date}."
        ),
        "application_status_followup": (
            "I am following up on the application I submitted on {start_date} "
            "regarding {main_subject}. I remain interested in the opportunity and "
            "would appreciate an update on its status by {end_date}."
        ),
        "interview_availability": (
            "Thank you for considering my profile regarding {main_subject}. "
            "I am available for an interview from {start_date} to {end_date}. "
            "Please confirm a suitable date, time, time zone and interview platform."
        ),
        "interview_confirmation": (
            "I am writing to confirm my availability for the interview scheduled "
            "between {start_date} and {end_date} regarding {main_subject}. "
            "Please share the meeting link, interview time and any preparation "
            "instructions, if not already provided."
        ),
        "interview_rescheduling": (
            "I request that my interview scheduled for {start_date} be rescheduled "
            "to a suitable time on or before {end_date} due to {main_subject}. "
            "I apologise for any inconvenience and remain interested in the position."
        ),
        "interview_cancellation": (
            "I am writing to cancel my interview scheduled between {start_date} "
            "and {end_date} due to {main_subject}. I apologise for any inconvenience "
            "and appreciate the time given to my application."
        ),
        "post_interview_thank_you": (
            "Thank you for the opportunity to interview on {start_date} regarding "
            "{main_subject}. The discussion strengthened my interest in the role. "
            "I appreciate your time and look forward to receiving an update by "
            "{end_date}."
        ),
        "interview_feedback": (
            "I am following up on the interview completed on {start_date} regarding "
            "{main_subject}. I would appreciate any available feedback and an update "
            "on the selection process by {end_date}."
        ),
        "next_round_followup": (
            "I am following up regarding the next interview stage for "
            "{main_subject}. I completed the previous stage on {start_date} and "
            "would appreciate confirmation of the next steps by {end_date}."
        ),
        "salary_discussion": (
            "I would like to discuss the compensation associated with "
            "{main_subject}. Based on the conversation held on {start_date}, "
            "I request clarification regarding the fixed pay, variable pay, "
            "benefits and proposed joining terms by {end_date}."
        ),
        "salary_negotiation": (
            "Thank you for sharing the compensation details on {start_date}. "
            "I request a review of the offer based on {main_subject}. I would "
            "appreciate the opportunity to discuss a revised compensation package "
            "before {end_date}."
        ),
        "document_submission": (
            "I am submitting the requested candidate documents concerning "
            "{main_subject}. These documents cover the required period from "
            "{start_date} to {end_date}. Please confirm their receipt and let me "
            "know if any additional documentation is required."
        ),
        "background_verification": (
            "Please find the information and documents required for background "
            "verification concerning {main_subject}. The relevant verification "
            "period is from {start_date} to {end_date}. Please let me know if any "
            "additional clarification is needed."
        ),
        "reference_submission": (
            "I am submitting the requested professional references regarding "
            "{main_subject}. They may be contacted between {start_date} and "
            "{end_date}. Please let me know if further information is required."
        ),
        "offer_letter_followup": (
            "I am following up on the offer discussed on {start_date} regarding "
            "{main_subject}. I remain interested in joining the organisation and "
            "would appreciate receiving the written offer and employment terms "
            "by {end_date}."
        ),
        "offer_acceptance": (
            "I am pleased to formally accept the employment offer concerning "
            "{main_subject}. I received the offer on {start_date} and confirm my "
            "availability to join on {end_date}, subject to the agreed terms."
        ),
        "offer_rejection": (
            "Thank you for the employment offer received on {start_date} regarding "
            "{main_subject}. After careful consideration, I will not be able to "
            "accept the offer by {end_date}. I appreciate the opportunity and the "
            "time invested in the selection process."
        ),
        "joining_confirmation": (
            "I am writing to confirm my joining date of {start_date} regarding "
            "{main_subject}. I will complete the required documentation and "
            "onboarding formalities by {end_date}. Please share any remaining "
            "instructions."
        ),
        "joining_extension": (
            "I request an extension of my joining date from {start_date} to "
            "{end_date} due to {main_subject}. I remain committed to joining the "
            "organisation and apologise for any inconvenience caused."
        ),
        "application_withdrawal": (
            "I am writing to withdraw my application submitted on {start_date} "
            "regarding {main_subject}, effective by {end_date}. I appreciate the "
            "time and consideration given to my application."
        ),
        "referral_request": (
            "I am seeking a professional referral regarding {main_subject}. "
            "The application remains open from {start_date} to {end_date}. "
            "If you find my experience suitable, I would appreciate your referral."
        ),
    },

    "Banking & Finance": {
        "bank_service": (
            "I request assistance with a banking matter concerning {main_subject}. "
            "The issue relates to the period from {start_date} to {end_date}. "
            "Kindly review it and provide the necessary resolution."
        ),
        "account_statement": (
            "Please provide my account statement for the period from {start_date} "
            "to {end_date}. The statement is required for {main_subject}. "
            "Kindly send it through the registered or authorised channel."
        ),
        "payment_issue": (
            "I am reporting a payment issue that occurred between {start_date} and "
            "{end_date}. The relevant details are {main_subject}. Kindly investigate "
            "the transaction and provide a resolution."
        ),
        "payment_confirmation": (
            "I am writing to confirm the payment made on {start_date} concerning "
            "{main_subject}. Kindly verify its receipt and provide confirmation "
            "by {end_date}."
        ),
        "refund_request": (
            "I request a refund regarding {main_subject}. The relevant transaction "
            "was initiated on {start_date}, and I request that the refund be "
            "processed by {end_date}. Please confirm the applicable procedure."
        ),
        "refund_followup": (
            "I am following up on the refund requested on {start_date} concerning "
            "{main_subject}. As it has not yet been received, please provide its "
            "current status and expected completion date by {end_date}."
        ),
        "transaction_dispute": (
            "I wish to dispute a transaction recorded between {start_date} and "
            "{end_date}. The reason for the dispute is {main_subject}. Kindly "
            "investigate it and temporarily secure the account if necessary."
        ),
        "unauthorized_transaction": (
            "I am reporting an unauthorised transaction identified on {start_date}. "
            "The transaction details are {main_subject}. Please block further "
            "unauthorised activity, investigate the matter and provide a resolution "
            "by {end_date}."
        ),
        "failed_transaction": (
            "I am reporting a failed transaction initiated on {start_date} "
            "concerning {main_subject}. The amount has not been correctly processed "
            "or reversed. Kindly investigate and resolve the issue by {end_date}."
        ),
        "loan_application": (
            "I wish to apply for a loan concerning {main_subject}. I require the "
            "facility from {start_date} and would appreciate a decision or update "
            "by {end_date}. Please share the eligibility and documentation requirements."
        ),
        "loan_status_followup": (
            "I am following up on my loan application submitted on {start_date} "
            "concerning {main_subject}. Please provide its current status and any "
            "remaining requirements by {end_date}."
        ),
        "insurance_claim": (
            "I am submitting an insurance claim concerning {main_subject}. "
            "The relevant incident or coverage period is from {start_date} to "
            "{end_date}. Kindly acknowledge the claim and provide the next steps."
        ),
        "invoice_submission": (
            "Please find the invoice concerning {main_subject} for the period from "
            "{start_date} to {end_date}. Kindly acknowledge receipt and process it "
            "according to the agreed payment terms."
        ),
        "payment_reminder": (
            "This is a reminder regarding the outstanding payment for "
            "{main_subject}. The payment was due on {start_date}. Kindly complete "
            "the payment or provide a status update by {end_date}."
        ),
    },

    "Customer & Product Service": {
        "product_enquiry": (
            "I would like information about {main_subject}. Please provide the "
            "product specifications, price, availability, warranty and delivery "
            "details for the period from {start_date} to {end_date}."
        ),
        "service_enquiry": (
            "I am seeking information regarding {main_subject}. Please share the "
            "service scope, pricing, availability and applicable terms between "
            "{start_date} and {end_date}."
        ),
        "technical_support": (
            "I require technical assistance regarding {main_subject}. The issue "
            "started on {start_date} and remains unresolved as of {end_date}. "
            "Kindly investigate and provide troubleshooting support."
        ),
        "product_complaint": (
            "I wish to report a problem with a product concerning {main_subject}. "
            "The issue was identified on {start_date} and continued until "
            "{end_date}. Kindly inspect the matter and provide an appropriate solution."
        ),
        "service_complaint": (
            "I wish to complain about a service issue concerning {main_subject}. "
            "The issue occurred between {start_date} and {end_date}. Please review "
            "the matter and provide an appropriate resolution."
        ),
        "complaint_escalation": (
            "I am escalating an unresolved complaint initially reported on "
            "{start_date} concerning {main_subject}. As it remains unresolved, "
            "I request senior review and a final response by {end_date}."
        ),
        "replacement_request": (
            "I request replacement of the product concerning {main_subject}. "
            "The issue was reported on {start_date}. Kindly confirm replacement "
            "eligibility and complete the process by {end_date}."
        ),
        "cancellation_request": (
            "I request cancellation of the order or service concerning "
            "{main_subject}, effective from {start_date}. Please confirm the "
            "cancellation and any applicable refund by {end_date}."
        ),
        "warranty_request": (
            "I request warranty assistance concerning {main_subject}. The issue "
            "was identified on {start_date}, and I request service or replacement "
            "by {end_date}. Please confirm the required procedure."
        ),
        "delivery_delay": (
            "I am reporting a delay in delivery concerning {main_subject}. "
            "The expected delivery date was {start_date}, but the order remains "
            "pending. Kindly provide an updated status and deliver it by {end_date}."
        ),
        "order_status": (
            "I request an update on the order concerning {main_subject}. "
            "It was placed on {start_date}, and the expected completion or delivery "
            "date is {end_date}. Please provide its current status."
        ),
        "service_feedback": (
            "I would like to share feedback regarding {main_subject} based on my "
            "experience from {start_date} to {end_date}. I hope this feedback helps "
            "improve the quality of the product or service."
        ),
    },

    "Study & University": {
        "admission_enquiry": (
            "I am seeking admission information concerning {main_subject}. "
            "Please provide the eligibility criteria, required documents, fees and "
            "application process for the period from {start_date} to {end_date}."
        ),
        "exam": (
            "I am writing regarding an examination matter concerning "
            "{main_subject}. The relevant examination period is from {start_date} "
            "to {end_date}. Kindly review my request and provide guidance."
        ),
        "exam_leave": (
            "I request examination leave from {start_date} to {end_date} due to "
            "{main_subject}. I will submit the examination schedule or supporting "
            "documents if required."
        ),
        "exam_schedule": (
            "I request clarification regarding the examination schedule concerning "
            "{main_subject}. Please confirm the dates, timings and venue for the "
            "period from {start_date} to {end_date}."
        ),
        "project_submission": (
            "I am submitting my academic project concerning {main_subject}. "
            "The project was completed during the period from {start_date} to "
            "{end_date}. Kindly acknowledge its receipt and advise if any correction "
            "or additional document is required."
        ),
        "project_guide_followup": (
            "I am following up regarding project guidance for {main_subject}. "
            "The work began on {start_date}, and the expected completion date is "
            "{end_date}. Kindly provide your review and further instructions."
        ),
        "practical_rescheduling": (
            "I request that my practical examination scheduled on {start_date} "
            "be rescheduled on or before {end_date} due to {main_subject}. "
            "Kindly consider my request and provide an alternative schedule."
        ),
        "viva_rescheduling": (
            "I request that my viva-voce scheduled on {start_date} be rescheduled "
            "on or before {end_date} due to {main_subject}. I will remain available "
            "on the alternative date approved by the department."
        ),
        "assignment_submission": (
            "I am submitting my assignment concerning {main_subject}. "
            "The assignment period is from {start_date} to {end_date}. "
            "Kindly acknowledge its submission."
        ),
        "assignment_extension": (
            "I request an extension of the assignment deadline from {start_date} "
            "to {end_date} due to {main_subject}. I will complete and submit the "
            "work within the requested extended period."
        ),
        "fees_payment": (
            "I am writing regarding a fee-related matter concerning {main_subject}. "
            "The relevant payment period is from {start_date} to {end_date}. "
            "Kindly review the details and provide the necessary resolution."
        ),
        "attendance_correction": (
            "I request correction of my academic attendance for the period from "
            "{start_date} to {end_date}. The reason for the requested correction "
            "is {main_subject}. Kindly verify the records and update them."
        ),
        "academic_complaint": (
            "I wish to submit a formal academic complaint concerning {main_subject}. "
            "The matter occurred between {start_date} and {end_date}. I request "
            "a fair review and an appropriate resolution."
        ),
        "certificate": (
            "I request the issuance of an academic certificate concerning "
            "{main_subject}. The relevant academic period is from {start_date} "
            "to {end_date}. Kindly inform me about the required procedure."
        ),
        "scholarship": (
            "I am applying for a scholarship concerning {main_subject}. "
            "The application period is from {start_date} to {end_date}. "
            "Kindly consider my application and attached supporting documents."
        ),
        "recommendation_letter": (
            "I request a letter of recommendation concerning {main_subject}. "
            "It is required for an application open from {start_date} to "
            "{end_date}. I would appreciate your support and recommendation."
        ),
    },

    "Business & Corporate": {
        "business_proposal": (
            "I am submitting a business proposal concerning {main_subject}. "
            "The proposed engagement would run from {start_date} to {end_date}. "
            "I would appreciate an opportunity to discuss the scope, commercial "
            "terms and expected outcomes."
        ),
        "partnership_proposal": (
            "I would like to propose a business partnership concerning "
            "{main_subject}. The proposed collaboration period is from "
            "{start_date} to {end_date}. I believe this partnership can create "
            "mutual value and welcome an opportunity to discuss it."
        ),
        "project_update": (
            "This message provides an update regarding {main_subject}. "
            "The reporting period is from {start_date} to {end_date}. "
            "Please review the current progress, pending activities and next steps."
        ),
        "project_delay": (
            "I am writing to notify you of a delay concerning {main_subject}. "
            "The original schedule began on {start_date}, and the revised expected "
            "completion date is {end_date}. We apologise for the inconvenience and "
            "are taking steps to complete the pending work."
        ),
        "business_followup": (
            "I am following up on our business discussion held on {start_date} "
            "regarding {main_subject}. I would appreciate your response or proposed "
            "next steps by {end_date}."
        ),
        "quotation_request": (
            "Please provide a quotation concerning {main_subject}. The required "
            "service or supply period is from {start_date} to {end_date}. "
            "Kindly include pricing, taxes, delivery terms and validity."
        ),
        "quotation_submission": (
            "Please find our quotation concerning {main_subject}. The quotation "
            "covers the period from {start_date} to {end_date}. We would be pleased "
            "to clarify the scope, pricing or commercial terms."
        ),
        "contract_discussion": (
            "I would like to discuss the contractual terms concerning "
            "{main_subject}. The proposed agreement period is from {start_date} "
            "to {end_date}. Kindly review the scope, responsibilities, payment "
            "terms and termination conditions."
        ),
        "vendor_communication": (
            "I am writing regarding vendor services concerning {main_subject}. "
            "The applicable period is from {start_date} to {end_date}. "
            "Please review the requirements and confirm your ability to proceed."
        ),
        "mou_request": (
            "I propose establishing a Memorandum of Understanding concerning "
            "{main_subject}. The proposed collaboration period is from "
            "{start_date} to {end_date}. Kindly review the proposal and confirm "
            "a suitable time for discussion."
        ),
    },

    "Meeting & Administration": {
        "meeting_request": (
            "I request a meeting regarding {main_subject}. Please schedule it at "
            "a mutually convenient time between {start_date} and {end_date}. "
            "The meeting will help us review the matter and agree on the next steps."
        ),
        "meeting_invitation": (
            "You are invited to attend a meeting regarding {main_subject}, "
            "scheduled between {start_date} and {end_date}. Please confirm your "
            "availability and review any shared agenda or supporting documents."
        ),
        "meeting_confirmation": (
            "This message confirms the meeting regarding {main_subject}. "
            "The confirmed meeting period is between {start_date} and {end_date}. "
            "Please join using the shared meeting details."
        ),
        "meeting_rescheduling": (
            "I request that the meeting scheduled on {start_date} regarding "
            "{main_subject} be rescheduled to a suitable time on or before "
            "{end_date}. I apologise for any inconvenience."
        ),
        "meeting_cancellation": (
            "This is to inform you that the meeting scheduled between {start_date} "
            "and {end_date} regarding {main_subject} has been cancelled. "
            "A revised schedule will be shared if required."
        ),
        "meeting_agenda": (
            "The meeting scheduled between {start_date} and {end_date} will discuss "
            "{main_subject}. Participants are requested to review the relevant "
            "documents and prepare any updates or decisions required."
        ),
        "mom_internal": (
            "This document records the minutes of the internal meeting conducted "
            "between {start_date} and {end_date} regarding {main_subject}. "
            "It summarises the discussion, decisions, responsibilities and deadlines."
        ),
        "mom_client": (
            "This document records the minutes of the client meeting conducted "
            "between {start_date} and {end_date} regarding {main_subject}. "
            "It summarises the agreed requirements, decisions and action items."
        ),
        "mom_project": (
            "This document records the project meeting held between {start_date} "
            "and {end_date} regarding {main_subject}. It summarises progress, "
            "issues, decisions, owners and delivery deadlines."
        ),
        "mom_review": (
            "This document records the review meeting held between {start_date} "
            "and {end_date} regarding {main_subject}. It captures the review "
            "observations, decisions, corrective actions and responsible persons."
        ),
        "action_item_summary": (
            "The following action items arise from the discussion held between "
            "{start_date} and {end_date} regarding {main_subject}. Each assigned "
            "owner should complete the relevant task within the agreed deadline."
        ),
        "meeting_followup": (
            "I am following up on the meeting held on {start_date} regarding "
            "{main_subject}. Kindly provide the pending updates and complete the "
            "agreed actions by {end_date}."
        ),
        "general_application": (
            "I am submitting this application regarding {main_subject}. "
            "The request applies from {start_date} to {end_date}. Kindly review "
            "the information and provide the required approval or response."
        ),
        "permission_application": (
            "I request permission regarding {main_subject} for the period from "
            "{start_date} to {end_date}. I will follow all applicable instructions "
            "and conditions if permission is granted."
        ),
        "complaint_application": (
            "I am submitting a formal complaint regarding {main_subject}. "
            "The issue occurred between {start_date} and {end_date}. Kindly "
            "investigate the matter and provide an appropriate resolution."
        ),
    },

    "General": {
        "help": (
            "I am writing to request assistance regarding {main_subject}. "
            "The matter concerns the period from {start_date} to {end_date}. "
            "Kindly review my request and provide the necessary guidance."
        ),
        "complaint": (
            "I wish to submit a complaint regarding {main_subject}. The issue "
            "occurred between {start_date} and {end_date}. Kindly investigate "
            "the matter and provide an appropriate resolution."
        ),
        "request": (
            "I am writing to submit a request regarding {main_subject}. "
            "The request applies from {start_date} to {end_date}. Kindly consider "
            "it and provide your approval or response."
        ),
        "enquiry": (
            "I am seeking information regarding {main_subject}. Please provide "
            "the relevant details for the period from {start_date} to {end_date}."
        ),
        "followup": (
            "I am following up regarding {main_subject}. The matter has been "
            "pending since {start_date}, and I would appreciate an update by "
            "{end_date}."
        ),
        "invitation": (
            "You are cordially invited regarding {main_subject}. The event or "
            "activity is scheduled between {start_date} and {end_date}. "
            "Please confirm your availability."
        ),
        "thank_you": (
            "I would like to express my sincere appreciation regarding "
            "{main_subject}. Your support between {start_date} and {end_date} "
            "was valuable and greatly appreciated."
        ),
        "apology": (
            "I sincerely apologise regarding {main_subject}. The matter occurred "
            "between {start_date} and {end_date}. I accept responsibility and "
            "will take appropriate steps to prevent it from happening again."
        ),
        "feedback": (
            "I would like to share my feedback regarding {main_subject}, based on "
            "my experience from {start_date} to {end_date}. I hope this feedback "
            "will support future improvement."
        ),
    },
}

DEFAULT_DESCRIPTION = (
    "I am writing regarding {main_subject}. The matter concerns the period "
    "from {start_date} to {end_date}. Kindly review the information and provide "
    "the necessary guidance, approval or resolution. Please let me know if any "
    "additional details or supporting documents are required from my side."
)

def db_connection():
    connection = sqlite3.connect(DB_PATH, timeout=30)
    connection.row_factory = sqlite3.Row
    connection.execute("PRAGMA foreign_keys = ON")
    connection.execute("PRAGMA busy_timeout = 30000")
    connection.execute("PRAGMA journal_mode = WAL")
    connection.execute("PRAGMA synchronous = NORMAL")
    return connection


def init_db():
    with db_connection() as db:
        db.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            email TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            designation TEXT DEFAULT '',
            organization TEXT DEFAULT '',
            role TEXT NOT NULL DEFAULT 'user',
            created_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS writings (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            writing_type TEXT NOT NULL,
            category_group TEXT NOT NULL,
            category TEXT NOT NULL,
            recipient_name TEXT DEFAULT '',
            recipient_email TEXT DEFAULT '',
            subject TEXT NOT NULL,
            body TEXT NOT NULL,
            attachment TEXT DEFAULT '',
            attachment_original_name TEXT DEFAULT '',
            signature_type TEXT DEFAULT 'none',
            signature_file TEXT DEFAULT '',
            language TEXT DEFAULT 'en',
            download_count INTEGER NOT NULL DEFAULT 0,
            print_count INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        );

        CREATE INDEX IF NOT EXISTS idx_writings_user ON writings(user_id);
        CREATE INDEX IF NOT EXISTS idx_writings_created ON writings(created_at);
        """)
    create_default_admin()


def create_default_admin():
    email = os.getenv("ADMIN_EMAIL", "admin@rppssoftware.com").strip().lower()
    password = os.getenv("ADMIN_PASSWORD", "Admin@123")
    with db_connection() as db:
        exists = db.execute("SELECT id FROM users WHERE email = ?", (email,)).fetchone()
        if not exists:
            db.execute(
                """INSERT INTO users
                   (id, name, email, password_hash, designation, organization, role, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, 'admin', ?)""",
                (
                    str(uuid.uuid4()),
                    "RPPS Administrator",
                    email,
                    generate_password_hash(password),
                    "Administrator",
                    "RPPS Software",
                    datetime.now().isoformat(timespec="seconds"),
                ),
            )
            db.commit()


def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if "user_id" not in session:
            flash("Please log in to continue.", "warning")
            return redirect(url_for("login"))
        return view(*args, **kwargs)

    return wrapped


def admin_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if session.get("role") != "admin":
            flash("Administrator access is required.", "danger")
            return redirect(url_for("dashboard"))
        return view(*args, **kwargs)

    return wrapped


def allowed_file(filename, image_only=False):
    if "." not in filename:
        return False
    extension = filename.rsplit(".", 1)[1].lower()
    return extension in (IMAGE_EXTENSIONS if image_only else ALLOWED_EXTENSIONS)


def clean_sentence(text):
    text = re.sub(r"\s+", " ", (text or "").strip())
    replacements = {
        r"\bi\b": "I",
        r"\bpls\b": "please",
        r"\bplz\b": "please",
        r"\bu\b": "you",
        r"\bur\b": "your",
        r"\bcoz\b": "because",
        r"\bcant\b": "cannot",
        r"\bwont\b": "will not",
        r"\bdidnt\b": "did not",
        r"\bdoesnt\b": "does not",
        r"\bim\b": "I am",
    }
    for pattern, value in replacements.items():
        text = re.sub(pattern, value, text, flags=re.IGNORECASE)
    if text and text[-1] not in ".!?":
        text += "."
    return text[:1].upper() + text[1:] if text else ""


def improve_subject(subject, category_label):
    subject = clean_sentence(subject).rstrip(".")
    return (
        subject
        if len(subject.split()) >= 3
        else (f"{category_label}: {subject}" if subject else category_label)
    )


def build_standard_body(writing_type, category_label, recipient_name, details, user):
    details = clean_sentence(details)
    recipient_name = recipient_name.strip() or "Sir/Madam"
    informal = writing_type == "informal_email"
    greeting = f"Hi {recipient_name}," if informal else f"Dear {recipient_name},"
    closing = (
        "Best regards,"
        if informal
        else (
            "Yours sincerely,"
            if writing_type in {"formal_letter", "application"}
            else "Kind regards,"
        )
    )

    opener = f"I am writing regarding {category_label.lower()}."
    request_line = "I kindly request you to review the matter and provide the necessary assistance or resolution."

    signature = [user["name"]]
    if user["designation"]:
        signature.append(user["designation"])
    if user["organization"]:
        signature.append(user["organization"])
    signature.append(user["email"])

    return "\n\n".join(
        [
            greeting,
            opener,
            details,
            request_line,
            "Thank you for your time and consideration.",
            closing,
            "\n".join(signature),
        ]
    )


def build_mom_body(subject, details, user):
    today = datetime.now().strftime("%d %B %Y")
    return "\n\n".join(
        [
            "MINUTES OF MEETING",
            f"Meeting Title: {subject}",
            f"Date: {today}",
            f"Prepared by: {user['name']}",
            "Attendees:\n[Add attendee names and roles]",
            "Agenda:\n[Add the meeting agenda]",
            f"Discussion Summary:\n{clean_sentence(details)}",
            "Decisions Taken:\n1. [Add decision]\n2. [Add decision]",
            "Action Items:\n1. [Task] — [Owner] — [Due date]\n2. [Task] — [Owner] — [Due date]",
            "Next Meeting:\n[Add date, time and purpose]",
            f"Prepared and recorded by:\n{user['name']}\n{user['organization'] or 'RPPS Writing Assistant'}",
        ]
    )


def translate_text(text, target_language):
    if not text or target_language == "en" or GoogleTranslator is None:
        return text
    try:
        output = []
        for part in text.split("\n\n"):
            output.append(
                GoogleTranslator(source="auto", target=target_language).translate(part)
                if part.strip()
                else ""
            )
        return "\n\n".join(output)
    except Exception:
        return text


def save_signature(data_url, uploaded):
    if uploaded and uploaded.filename:
        if not allowed_file(uploaded.filename, image_only=True):
            raise ValueError("Signature must be a PNG, JPG or JPEG image.")
        original = secure_filename(uploaded.filename)
        filename = f"signature_{uuid.uuid4().hex}_{original}"
        uploaded.save(UPLOAD_DIR / filename)
        return "uploaded", filename
    if data_url and data_url.startswith("data:image"):
        _, encoded = data_url.split(",", 1)
        filename = f"signature_{uuid.uuid4().hex}.png"
        (UPLOAD_DIR / filename).write_bytes(base64.b64decode(encoded))
        return "drawn", filename
    return "none", ""


def current_user():
    with db_connection() as db:
        return db.execute(
            "SELECT * FROM users WHERE id = ?", (session["user_id"],)
        ).fetchone()


def writing_by_id(writing_id):
    with db_connection() as db:
        query = "SELECT * FROM writings WHERE id = ?"
        params = [writing_id]
        if session.get("role") != "admin":
            query += " AND user_id = ?"
            params.append(session["user_id"])
        return db.execute(query, params).fetchone()


init_db()

def format_date(date_value):
    """
    Converts 2026-07-15 into 15 July 2026.
    Returns 'Not specified' when the date is empty.
    """
    if not date_value:
        return "Not specified"

    try:
        parsed_date = datetime.strptime(date_value, "%Y-%m-%d")
        return parsed_date.strftime("%d %B %Y")
    except (ValueError, TypeError):
        return str(date_value)


def generate_predefined_content(
    area,
    category_key,
    start_date,
    end_date,
    main_subject
):
    title_template = (
        PREDEFINED_TITLES
        .get(area, {})
        .get(category_key)
    )

    if not title_template:
        category_label = (
            CATEGORY_TREE
            .get(area, {})
            .get(category_key, category_key.replace("_", " ").title())
        )
        title_template = f"Regarding {category_label}"

    description_template = (
        PREDEFINED_DESCRIPTIONS
        .get(area, {})
        .get(category_key, DEFAULT_DESCRIPTION)
    )

    values = {
        "start_date": format_date(start_date),
        "end_date": format_date(end_date),
        "main_subject": main_subject.strip() or "the stated matter",
    }

    return {
        "title": title_template,
        "description": description_template.format(**values),
    }


@app.route("/")
def index():
    return render_template("index.html")


def render_info_page(page_name):
    return render_template("info_page.html", page=INFO_PAGES[page_name])


@app.route("/about")
def about():
    return render_info_page("about")


@app.route("/services")
def services():
    return render_info_page("services")


@app.route("/contact")
def contact():
    return render_info_page("contact")


@app.route("/privacy")
def privacy():
    return render_info_page("privacy")


@app.route("/terms")
def terms():
    return render_info_page("terms")


@app.route("/help")
def help_support():
    return render_info_page("help")


@app.route("/accessibility")
def accessibility():
    return render_info_page("accessibility")


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        if not name or not email or len(password) < 6:
            flash(
                "Name, email and a password of at least six characters are required.",
                "danger",
            )
            return render_template("register.html")
        try:
            with db_connection() as db:
                db.execute(
                    """INSERT INTO users
                       (id, name, email, password_hash, designation, organization, role, created_at)
                       VALUES (?, ?, ?, ?, ?, ?, 'user', ?)""",
                    (
                        str(uuid.uuid4()),
                        name,
                        email,
                        generate_password_hash(password),
                        request.form.get("designation", "").strip(),
                        request.form.get("organization", "").strip(),
                        datetime.now().isoformat(timespec="seconds"),
                    ),
                )
                db.commit()
            flash("Registration completed. Please log in.", "success")
            return redirect(url_for("login"))
        except sqlite3.IntegrityError:
            flash("An account with this email already exists.", "warning")
    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        with db_connection() as db:
            user = db.execute(
                "SELECT * FROM users WHERE email = ?", (email,)
            ).fetchone()
        if user and check_password_hash(
            user["password_hash"], request.form.get("password", "")
        ):
            session.clear()
            session.permanent = True
            session["user_id"] = user["id"]
            session["user_name"] = user["name"]
            session["role"] = user["role"]
            return redirect(
                url_for("admin_dashboard" if user["role"] == "admin" else "dashboard")
            )
        flash("Invalid email or password.", "danger")
    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    flash("You have been logged out.", "info")
    return redirect(url_for("index"))


@app.route("/dashboard")
@login_required
def dashboard():
    if session.get("role") == "admin":
        return redirect(url_for("admin_dashboard"))
    with db_connection() as db:
        stats = db.execute(
            """SELECT COUNT(*) writing_count,
                      COALESCE(SUM(download_count), 0) download_count,
                      COALESCE(SUM(print_count), 0) print_count
               FROM writings WHERE user_id = ?""",
            (session["user_id"],),
        ).fetchone()
        records = db.execute(
            "SELECT * FROM writings WHERE user_id = ? ORDER BY created_at DESC LIMIT 8",
            (session["user_id"],),
        ).fetchall()
    return render_template("dashboard.html", records=records, stats=stats)


@app.route("/admin")
@login_required
@admin_required
def admin_dashboard():
    with db_connection() as db:
        totals = db.execute(
            """SELECT
                (SELECT COUNT(*) FROM users WHERE role='user') total_users,
                (SELECT COUNT(*) FROM writings) total_writings,
                (SELECT COALESCE(SUM(download_count),0) FROM writings) total_downloads,
                (SELECT COALESCE(SUM(print_count),0) FROM writings) total_prints"""
        ).fetchone()
        users = db.execute(
            """SELECT u.id, u.name, u.email, u.designation, u.organization, u.created_at,
                      COUNT(w.id) writing_count,
                      COALESCE(SUM(w.download_count),0) download_count,
                      COALESCE(SUM(w.print_count),0) print_count
               FROM users u
               LEFT JOIN writings w ON w.user_id=u.id
               WHERE u.role='user'
               GROUP BY u.id
               ORDER BY u.created_at DESC"""
        ).fetchall()
    return render_template("admin_dashboard.html", totals=totals, users=users)


@app.route("/admin/export/<file_type>")
@login_required
@admin_required
def admin_export(file_type):
    with db_connection() as db:
        users = pd.read_sql_query(
            "SELECT id,name,email,designation,organization,role,created_at FROM users",
            db,
        )
        writings = pd.read_sql_query("SELECT * FROM writings", db)

    output = io.BytesIO()
    if file_type == "xlsx":
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            users.to_excel(writer, sheet_name="Users", index=False)
            writings.to_excel(writer, sheet_name="Writings", index=False)
        output.seek(0)
        return send_file(
            output, as_attachment=True, download_name="rpps_admin_report.xlsx"
        )
    if file_type == "csv":
        output.write(users.to_csv(index=False).encode("utf-8"))
        output.seek(0)
        return send_file(
            output,
            as_attachment=True,
            download_name="rpps_users.csv",
            mimetype="text/csv",
        )
    flash("Unsupported export type.", "warning")
    return redirect(url_for("admin_dashboard"))


@app.route("/compose", methods=["GET", "POST"])
@login_required
def compose():
    if request.method == "POST":
        writing_type = request.form.get("writing_type", "formal_email")
        group = request.form.get("category_group", "General")
        category = request.form.get("category", "help")
        category_label = next(
            (items[category] for items in CATEGORY_TREE.values() if category in items),
            "General Request",
        )
        recipient_name = request.form.get("recipient_name", "").strip()
        recipient_email = request.form.get("recipient_email", "").strip()
        language = request.form.get("language", "en")
        subject = improve_subject(request.form.get("subject", ""), category_label)
        details = request.form.get("details", "")

        attachment = attachment_original = ""
        uploaded = request.files.get("attachment")
        if uploaded and uploaded.filename:
            if not allowed_file(uploaded.filename):
                flash("Unsupported attachment type.", "danger")
                return redirect(url_for("compose"))
            attachment_original = secure_filename(uploaded.filename)
            attachment = f"{uuid.uuid4().hex}_{attachment_original}"
            uploaded.save(UPLOAD_DIR / attachment)

        try:
            signature_type, signature_file = save_signature(
                request.form.get("signature_data", ""),
                request.files.get("signature_upload"),
            )
        except ValueError as exc:
            flash(str(exc), "danger")
            return redirect(url_for("compose"))

        user = current_user()
        body = (
            build_mom_body(subject, details, user)
            if writing_type == "mom"
            else build_standard_body(
                writing_type, category_label, recipient_name, details, user
            )
        )
        subject = translate_text(subject, language)
        body = translate_text(body, language)
        writing_id = str(uuid.uuid4())

        with db_connection() as db:
            db.execute(
                """INSERT INTO writings
                   (id,user_id,writing_type,category_group,category,recipient_name,
                    recipient_email,subject,body,attachment,attachment_original_name,
                    signature_type,signature_file,language,created_at)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (
                    writing_id,
                    session["user_id"],
                    writing_type,
                    group,
                    category,
                    recipient_name,
                    recipient_email,
                    subject,
                    body,
                    attachment,
                    attachment_original,
                    signature_type,
                    signature_file,
                    language,
                    datetime.now().isoformat(timespec="seconds"),
                ),
            )
            db.commit()
        return redirect(url_for("preview", writing_id=writing_id))

    return render_template(
        "compose.html",
        writing_types=WRITING_TYPES,
        category_tree=CATEGORY_TREE,
        languages=LANGUAGES,
    )


@app.route("/ocr", methods=["POST"])
@login_required
def ocr_image():
    image_file = request.files.get("handwriting_image")
    if not image_file or not image_file.filename:
        return jsonify({"error": "Select a handwritten image first."}), 400
    if not allowed_file(image_file.filename, image_only=True):
        return jsonify({"error": "Only PNG, JPG and JPEG images are supported."}), 400
    if pytesseract is None:
        return jsonify({"error": "pytesseract is not installed on the server."}), 503
    try:
        image = Image.open(image_file.stream).convert("RGB")
        language = request.form.get("ocr_language", "eng")
        text = pytesseract.image_to_string(image, lang=language).strip()
        return jsonify({"text": text})
    except Exception as exc:
        return (
            jsonify(
                {
                    "error": "OCR failed. Confirm that the Tesseract OCR system package and selected language pack are installed."
                }
            ),
            500,
        )


@app.route("/preview/<writing_id>")
@login_required
def preview(writing_id):
    writing = writing_by_id(writing_id)
    if not writing:
        flash("Writing not found.", "danger")
        return redirect(url_for("dashboard"))
    return render_template("preview.html", writing=writing)


@app.route("/history")
@login_required
def history():
    with db_connection() as db:
        records = db.execute(
            "SELECT * FROM writings WHERE user_id=? ORDER BY created_at DESC",
            (session["user_id"],),
        ).fetchall()
    return render_template("history.html", records=records)


@app.route("/count-print/<writing_id>", methods=["POST"])
@login_required
def count_print(writing_id):
    writing = writing_by_id(writing_id)
    if not writing:
        return jsonify({"error": "Not found"}), 404
    with db_connection() as db:
        db.execute(
            """UPDATE writings
               SET print_count=print_count+1,
                   download_count=download_count+1
               WHERE id=?""",
            (writing_id,),
        )
        db.commit()
    return jsonify({"success": True})


@app.route("/download/<writing_id>/<file_type>")
@login_required
def download(writing_id, file_type):
    writing = writing_by_id(writing_id)
    if not writing:
        flash("Writing not found.", "danger")
        return redirect(url_for("dashboard"))

    with db_connection() as db:
        db.execute(
            "UPDATE writings SET download_count=download_count+1 WHERE id=?",
            (writing_id,),
        )
        db.commit()

    subject, body = writing["subject"], writing["body"]
    safe_base = re.sub(r"[^A-Za-z0-9_-]+", "_", subject)[:60] or "writing"

    if file_type == "txt":
        return send_file(
            io.BytesIO(f"Subject: {subject}\n\n{body}".encode()),
            as_attachment=True,
            download_name=f"{safe_base}.txt",
            mimetype="text/plain",
        )

    if file_type == "docx":
        document = Document()
        document.add_heading(subject, level=1)
        for paragraph in body.split("\n\n"):
            document.add_paragraph(paragraph)
        if writing["signature_file"]:
            document.add_picture(str(UPLOAD_DIR / writing["signature_file"]))
        if writing["attachment_original_name"]:
            document.add_paragraph(f"Attachment: {writing['attachment_original_name']}")
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name=f"{safe_base}.docx")

    if file_type == "pdf":
        output = io.BytesIO()
        document = SimpleDocTemplate(
            output,
            pagesize=A4,
            rightMargin=20 * mm,
            leftMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
        )
        styles = getSampleStyleSheet()
        story = [Paragraph(subject, styles["Title"]), Spacer(1, 10)]
        for paragraph in body.split("\n\n"):
            escaped = (
                paragraph.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace("\n", "<br/>")
            )
            story.extend([Paragraph(escaped, styles["BodyText"]), Spacer(1, 8)])
        if writing["signature_file"]:
            story.extend(
                [
                    Spacer(1, 10),
                    PDFImage(
                        str(UPLOAD_DIR / writing["signature_file"]),
                        width=50 * mm,
                        height=20 * mm,
                    ),
                ]
            )
        if writing["attachment_original_name"]:
            story.extend(
                [
                    Spacer(1, 15),
                    Paragraph(
                        f"Attachment: {writing['attachment_original_name']}",
                        styles["BodyText"],
                    ),
                ]
            )
        document.build(story)
        output.seek(0)
        return send_file(
            output,
            as_attachment=True,
            download_name=f"{safe_base}.pdf",
            mimetype="application/pdf",
        )

    if file_type == "png":
        width, margin = 1400, 80
        font = ImageFont.load_default(size=25)
        title_font = ImageFont.load_default(size=34)
        lines = []
        for paragraph in body.splitlines():
            words, line = paragraph.split(), ""
            for word in words:
                test = f"{line} {word}".strip()
                if len(test) > 85:
                    lines.append(line)
                    line = word
                else:
                    line = test
            lines.append(line)
        height = max(1000, 240 + len(lines) * 38)
        image = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(image)
        draw.text((margin, margin), subject, fill="black", font=title_font)
        y = margin + 70
        for line in lines:
            draw.text((margin, y), line, fill="black", font=font)
            y += 38
        output = io.BytesIO()
        image.save(output, "PNG")
        output.seek(0)
        return send_file(
            output,
            as_attachment=True,
            download_name=f"{safe_base}.png",
            mimetype="image/png",
        )

    flash("Unsupported download type.", "warning")
    return redirect(url_for("preview", writing_id=writing_id))


@app.route("/uploads/<filename>")
@login_required
def uploaded_file(filename):
    safe_name = secure_filename(filename)
    file_path = UPLOAD_DIR / safe_name
    if not file_path.exists():
        flash("File not found.", "danger")
        return redirect(url_for("dashboard"))
    return send_file(file_path)

@app.post("/generate-predefined-content")
@login_required
def generate_content():
    data = request.get_json(silent=True) or request.form

    area = data.get("area", "")
    category_key = data.get("category", "")
    start_date = data.get("start_date", "")
    end_date = data.get("end_date", "")
    main_subject = data.get("main_subject", "")

    if area not in CATEGORY_TREE:
        return jsonify({
            "success": False,
            "error": "Please select a valid area."
        }), 400

    if category_key not in CATEGORY_TREE[area]:
        return jsonify({
            "success": False,
            "error": "Please select a valid category."
        }), 400

    content = generate_predefined_content(
        area=area,
        category_key=category_key,
        start_date=start_date,
        end_date=end_date,
        main_subject=main_subject
    )

    return jsonify({
        "success": True,
        "title": content["title"],
        "description": content["description"]
    })

@app.errorhandler(413)
def too_large(_):
    flash("File is too large. Maximum size is 10 MB.", "danger")
    return redirect(url_for("compose"))


if __name__ == "__main__":
    app.run(debug=True, threaded=True)
